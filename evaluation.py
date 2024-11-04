import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import os

class EvaluationWriter:
    def __init__(self):
        self.ratings = {
            "IDEAL": "✅ Complete and detailed response with evidence",
            "GOOD": "🟡 Adequate but could be improved",
            "BAD": "❌ Incomplete or inadequate response"
        }

    def write_formatted_results(self, results, questionnaire):
        """Write formatted results with evaluation buttons to both text and Word files"""
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H-%M-%S")
        txt_output = f"security_questionnaire_responses_{timestamp}.txt"
        docx_output = f"security_questionnaire_responses_{timestamp}.docx"
        
        # Create Streamlit interface
        st.title("Security Questionnaire Evaluation")
        
        # Initialize document
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        with open(txt_output, 'w', encoding='utf-8') as f:
            self._write_header(f, doc)
            
            for section, questions in results['answers'].items():
                section_name = questionnaire.sections.get(section, section)
                self._write_section(f, doc, section_name, questions, questionnaire)
                
        doc.save(docx_output)
        return txt_output, docx_output

    def _write_header(self, f, doc):
        # Reference to original header format
        """
        startLine: 226
        endLine: 231
        """

    def _write_section(self, f, doc, section_name, questions, questionnaire):
        st.header(section_name.upper())
        
        for q_key, answer in questions.items():
            full_question = questionnaire.questions[q_key]
            
            # Display question and answer
            st.subheader("Question:")
            st.write(full_question)
            st.subheader("Answer:")
            st.write(answer['result'] if isinstance(answer, dict) else answer)
            
            # Display source documents if they exist
            if 'sources' in answer:
                st.subheader("Source Documents:")
                for source in answer['sources']:
                    st.write(f"• {source}")
            
            # Add rating buttons
            col1, col2, col3 = st.columns(3)
            rating = None
            
            with col1:
                if st.button("✅ IDEAL", key=f"ideal_{q_key}"):
                    rating = "IDEAL"
            with col2:
                if st.button("🟡 GOOD", key=f"good_{q_key}"):
                    rating = "GOOD"
            with col3:
                if st.button("❌ BAD", key=f"bad_{q_key}"):
                    rating = "BAD"
            
            if rating:
                self._write_evaluation(f, doc, rating, answer)
            
            st.markdown("---")

    def _write_evaluation(self, f, doc, rating, answer):
        evaluation_text = f"📊 Evaluation:\nRating: {rating}\nJustification: {self.ratings[rating]}\n\n"
        f.write(evaluation_text)
        
        e_para = doc.add_paragraph()
        e_para.add_run("Evaluation:\n").bold = True
        e_para.add_run(f"Rating: {rating}\n")
        e_para.add_run(f"Justification: {self.ratings[rating]}") 