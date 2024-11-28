import streamlit as st
from vision import SecurityQuestionnaire, main as process_documents
import os
from datetime import datetime
from docx import Document
import glob
import time
import pandas as pd
from database import ResponseDatabase

# Initialize database at the top of the file
db = ResponseDatabase()

# Add this function after the imports and before apply_custom_styling()
def load_latest_docx():
    """Load the latest generated docx file"""
    try:
        docx_files = glob.glob("security_questionnaire_responses_*.docx")
        if not docx_files:
            return None
        latest_file = max(docx_files, key=os.path.getctime)
        return Document(latest_file)
    except Exception as e:
        st.error(f"Error loading document: {str(e)}")
        return None

# Custom theme and styling
def apply_custom_styling(theme=None):
    """Apply custom styling with HyperComply-like design"""
    st.markdown("""
        <style>
        /* Main layout */
        .main {
            background: #F9FAFB;
            padding: 2rem;
        }
        
        /* Custom header styling */
        .custom-header {
            background: linear-gradient(135deg, #1e40af 0%, #3b82f6 100%);
            padding: 2.5rem;
            border-radius: 12px;
            margin-bottom: 2rem;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        }
        
        .custom-header h1 {
            color: white !important;
            font-size: 2.5rem !important;
            font-weight: 700 !important;
            margin-bottom: 1rem !important;
            text-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
        }
        
        .custom-header p {
            color: rgba(255, 255, 255, 0.9) !important;
            font-size: 1.1rem !important;
            line-height: 1.5;
        }
        
        /* Radio buttons */
        .stRadio > div {
            display: flex;
            gap: 20px;
        }
        
        .stRadio [role="radiogroup"] {
            gap: 20px;
        }
        
        .stRadio label {
            background: white;
            padding: 8px 16px;
            border: 1px solid #E5E7EB;
            border-radius: 6px;
            color: #374151;
            font-weight: 500;
            min-width: 80px;
            text-align: center;
        }
        
        /* Text area */
        .stTextArea textarea {
            border: 1px solid #E5E7EB;
            border-radius: 6px;
            padding: 12px;
            font-size: 0.875rem;
            background: #F9FAFB;
        }
        
        /* Buttons */
        .stButton button {
            border: 1px solid #E5E7EB;
            background: white;
            color: #374151;
            font-weight: 500;
            padding: 8px 16px;
            border-radius: 6px;
            transition: all 0.2s ease;
            box-shadow: 0 1px 2px rgba(0, 0, 0, 0.05);
        }
        
        .stButton button:hover {
            transform: translateY(-1px);
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        }
        
        .stButton button[kind="primary"] {
            background: #059669;
            color: white;
            border: none;
        }
        
        /* Search bar */
        .stTextInput input {
            border: 1px solid #E5E7EB;
            border-radius: 6px;
            padding: 12px;
        }
        
        /* Button styling for Yes/No options */
        .stButton button[type="secondary"] {
            background: white;
            border: 1px solid #E5E7EB;
            color: #374151;
            min-width: 80px;
        }
        
        .stButton button[type="primary"] {
            background: #2563EB;
            color: white;
            border: none;
            min-width: 80px;
        }
        
        /* Answer blob styling */
        .answer-blob {
            background: #F3F4F6;
            border-radius: 12px;
            padding: 20px;
            margin: 15px 0;
            color: #374151;
            font-size: 1rem;
            line-height: 1.6;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
        }
        
        /* Enhanced Yes/No button styling */
        .stButton button {
            transition: all 0.2s ease;
            box-shadow: 0 1px 2px rgba(0, 0, 0, 0.05);
        }
        
        .stButton button:hover {
            transform: translateY(-1px);
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        }
        
        /* Enhanced answer blob styling */
        .answer-blob {
            background: #F3F4F6;
            border-radius: 12px;
            padding: 20px;
            margin: 15px 0;
            color: #374151;
            font-size: 1rem;
            line-height: 1.6;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
        }
        
        /* Action button styling */
        .action-buttons {
            margin-top: 15px;
            padding-top: 15px;
            border-top: 1px solid #E5E7EB;
        }
        </style>
    """, unsafe_allow_html=True)

def display_qa_section(question, answer, sources, qa_id):
    """Display a Q&A section with database integration"""
    with st.container():
        # Question container with white background and border
        st.markdown("""
            <div style='background: white; border: 1px solid #e5e7eb; border-radius: 8px; padding: 20px; margin-bottom: 20px;'>
        """, unsafe_allow_html=True)
        
        # Question header with icon and accuracy badge
        st.markdown(f"""
            <div style='display: flex; justify-content: space-between; align-items: center; margin-bottom: 15px;'>
                <div style='display: flex; align-items: center; gap: 10px;'>
                    <span style='color: #6B7280;'>⚙️</span>
                    <span style='color: #374151; font-weight: 500;'>{question}</span>
                </div>
                <div style='background: #EFF6FF; color: #2563EB; padding: 4px 12px; border-radius: 4px; font-size: 0.875rem;'>
                    Exact match
                </div>
            </div>
        """, unsafe_allow_html=True)
        
        # Get response history
        response_history = db.get_response_history(qa_id)
        
        if answer:
            st.markdown(f"""
                <div style='background: #F3F4F6; border-radius: 12px; padding: 20px; margin: 12px 0;'>
                    <div style='color: #374151; font-size: 1rem; line-height: 1.6; margin-bottom: 15px;'>
                        {answer}
                    </div>
                    <div style='border-top: 1px solid #E5E7EB; padding-top: 15px; margin-top: 15px;'>
                        <p style='color: #6B7280; font-size: 0.875rem; margin-bottom: 10px;'>Is this answer correct?</p>
                    </div>
                </div>
            """, unsafe_allow_html=True)
            
            # Yes/No buttons with database integration
            col1, col2, *rest = st.columns([1, 1, 4])
            with col1:
                if st.button("✓ Yes", key=f"yes_{qa_id}"):
                    response_id = db.save_response(
                        question_id=qa_id,
                        question=question,
                        answer=answer,
                        response="Yes",
                        sources=sources
                    )
                    st.success("Response saved!")
            
            with col2:
                if st.button("✕ No", key=f"no_{qa_id}"):
                    response_id = db.save_response(
                        question_id=qa_id,
                        question=question,
                        answer=answer,
                        response="No",
                        sources=sources
                    )
                    st.success("Response saved!")
            
            # Action buttons with database integration
            cols = st.columns([1, 1, 1, 3])
            with cols[0]:
                if st.button("✓ Apply", key=f"apply_{qa_id}", type="primary"):
                    if response_history:
                        db.update_response(response_history[0][0], "Applied")
                        st.success("Response applied!")
            
            with cols[1]:
                st.button("📎 Attach", key=f"attach_{qa_id}")
            
            with cols[2]:
                st.button("👤 Assign", key=f"assign_{qa_id}")
            
            # Show response history
            if response_history:
                with st.expander("View Response History"):
                    for resp in response_history:
                        st.markdown(f"""
                            **Response:** {resp[3]}  
                            **Date:** {resp[6]}  
                            **Status:** {resp[4]}
                            ---
                        """)
        
        st.markdown("</div>", unsafe_allow_html=True)

def show_progress_bar():
    progress_bar = st.progress(0)
    for i in range(100):
        time.sleep(0.01)
        progress_bar.progress(i + 1)
    return progress_bar

def copy_button(text, button_id):
    """Add a copy button for text content"""
    if st.button(f"📋 Copy", key=f"copy_{button_id}"):
        st.code(text)  # Shows copyable text in a code block
        st.success("Text copied to clipboard!")

def add_feedback_system(qa_id, response_id):
    """Add feedback buttons with database integration"""
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("👍 Helpful", key=f"helpful_{qa_id}"):
            db.save_feedback(response_id, "helpful")
            st.success("Thank you for your feedback!")
    with col2:
        if st.button("👎 Not Helpful", key=f"not_helpful_{qa_id}"):
            feedback = st.text_area("How can we improve?", key=f"feedback_{qa_id}")
            if st.button("Submit Feedback"):
                db.save_feedback(response_id, "not_helpful", feedback)
                st.success("Thank you for your feedback!")
    with col3:
        if st.button("⚠️ Report Issue", key=f"report_{qa_id}"):
            issue = st.text_area("Describe the issue:", key=f"issue_{qa_id}")
            if st.button("Submit Report"):
                db.save_feedback(response_id, "issue", issue)
                st.success("Issue reported!")

def add_theme_toggle():
    """Add theme toggle in sidebar"""
    with st.sidebar:
        theme = st.selectbox(
            "🎨 Choose Theme",
            ["Light", "Dark"],
            key="theme"
        )
        return theme

def add_export_options():
    """Add export options for the analysis"""
    with st.sidebar:
        st.markdown("### 📤 Export Options")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📄 Export PDF"):
                st.download_button(
                    label="Download PDF",
                    data=generate_pdf(),  # You'll need to implement this
                    file_name="analysis.pdf",
                    mime="application/pdf"
                )
        with col2:
            if st.button("📊 Export Excel"):
                st.download_button(
                    label="Download Excel",
                    data=generate_excel(),  # You'll need to implement this
                    file_name="analysis.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

def generate_pdf():
    """Generate PDF from analysis results"""
    try:
        # Get the latest docx file
        docx_files = glob.glob("security_questionnaire_responses_*.docx")
        if not docx_files:
            return None
        latest_file = max(docx_files, key=os.path.getctime)
        
        # Read the file as bytes
        with open(latest_file, 'rb') as f:
            return f.read()
    except Exception as e:
        st.error(f"Error generating PDF: {str(e)}")
        return None

def generate_excel():
    """Generate Excel from analysis results"""
    try:
        # Get the latest docx file
        docx_files = glob.glob("security_questionnaire_responses_*.docx")
        if not docx_files:
            return None
            
        doc = Document(max(docx_files, key=os.path.getctime))
        data = []
        
        for para in doc.paragraphs:
            if "Question:" in para.text:
                data.append({"Question": para.text.replace("Question:", "").strip()})
            elif "Answer:" in para.text and data:
                data[-1]["Answer"] = para.text.replace("Answer:", "").strip()
                
        df = pd.DataFrame(data)
        return df.to_excel(index=False).encode()
    except Exception as e:
        st.error(f"Error generating Excel: {str(e)}")
        return None

def add_knowledge_base():
    with st.sidebar:
        st.markdown("""
            <div class='knowledge-base'>
                <h3>Search the Knowledge Base</h3>
                <div class='search-results'>
                    <p>Results (20)</p>
                    <div style='margin-top: 1rem;'>
                        <p><strong>From past completed questionnaires</strong></p>
                        <!-- Add search results here -->
                    </div>
                </div>
            </div>
        """, unsafe_allow_html=True)

def main():
    st.set_page_config(
        page_title="Security Documentation Analyzer",
        page_icon="🔒",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Add theme toggle
    theme = add_theme_toggle()
    apply_custom_styling(theme)
    
    # Modern header with gradient
    st.markdown("""
        <div class='custom-header'>
            <h1 style='margin: 0; font-size: 3rem; font-weight: 800; color: white;'>
                Security Documentation Analyzer
            </h1>
            <p style='margin: 1rem 0 0 0; font-size: 1.25rem; color: rgba(255,255,255,0.9);'>
                Automate security questionnaires and accelerate your compliance process
            </p>
        </div>
    """, unsafe_allow_html=True)

    # Sidebar with improved styling
    with st.sidebar:
        st.markdown("""
            <div style='padding: 1rem; background: white; border-radius: 10px; box-shadow: 0 2px 4px rgba(0,0,0,0.05);'>
                <h2 style='color: #1e3c72; margin-bottom: 1rem;'>📤 Document Upload</h2>
            </div>
        """, unsafe_allow_html=True)
        
        uploaded_files = st.file_uploader(
            "Upload your security documentation",
            type=['pdf'],
            accept_multiple_files=True,
            help="Upload PDF files containing security documentation"
        )
        
        if uploaded_files:
            save_path = "docs"
            os.makedirs(save_path, exist_ok=True)
            
            st.markdown("### 📁 Uploaded Files")
            for file in uploaded_files:
                st.markdown(f"""
                    <div class='uploadedFile'>
                        <span style='color: #1e3c72;'>📄</span> {file.name}
                    </div>
                """, unsafe_allow_html=True)
                
                file_path = os.path.join(save_path, file.name)
                with open(file_path, "wb") as f:
                    f.write(file.getbuffer())
            
            if st.button("🔍 Analyze Documents", use_container_width=True):
                progress_bar = show_progress_bar()
                with st.spinner("📊 Processing documents..."):
                    process_documents()
                progress_bar.empty()
                st.success("✅ Analysis complete!")
                st.balloons()

    # Main content tabs
    tab1, tab2 = st.tabs(["📊 Analysis Results", "📝 Questions Database"])

    with tab1:
        # Add single search bar at the top
        search_query = st.text_input("🔍 Search questions and answers...", key="global_search").lower()
        
        doc = load_latest_docx()
        if doc:
            current_section = None
            qa_data = {"question": "", "answer": "", "sources": ""}
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    if "=" * 20 in text:
                        if current_section:
                            st.markdown("---")
                        current_section = text.replace("=", "").strip()
                        st.markdown(f"""
                            <div class='section-header'>
                                <h2 style='color: #1e3c72; margin: 0;'>{current_section}</h2>
                            </div>
                        """, unsafe_allow_html=True)
                    elif "Question:" in text:
                        if qa_data["question"]:
                            # Only display if matches search query
                            if not search_query or search_query in qa_data["question"].lower() or search_query in qa_data["answer"].lower():
                                display_qa_section(qa_data["question"], qa_data["answer"], qa_data["sources"], qa_data["question"])
                            qa_data = {"question": "", "answer": "", "sources": ""}
                        qa_data["question"] = text.replace("Question:", "").strip()
                    elif "Answer:" in text:
                        qa_data["answer"] = text.replace("Answer:", "").strip()
                    elif "Source Documents:" in text:
                        qa_data["sources"] = text.replace("Source Documents:", "").strip()
                        display_qa_section(qa_data["question"], qa_data["answer"], qa_data["sources"], qa_data["question"])
                        qa_data = {"question": "", "answer": "", "sources": ""}
        else:
            st.info("🔍 No analysis results available. Please upload documents and run the analysis.")

    with tab2:
        questionnaire = SecurityQuestionnaire()
        for section, questions in questionnaire.questions.items():
            with st.expander(f"📋 {section.replace('_', ' ').title()}", expanded=False):
                st.markdown("""
                    <div class='content-box'>
                """, unsafe_allow_html=True)
                for q_key, question in questions.items():
                    st.markdown(f"""
                        <div style='margin-bottom: 1rem;'>
                            <p style='color: #1e3c72; font-weight: 600; margin-bottom: 0.5rem;'>{q_key}</p>
                            <p style='color: #4a5568; font-style: italic;'>{question}</p>
                        </div>
                        <hr style='margin: 1rem 0; border-color: #e2e8f0;'>
                    """, unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)

    # Add export options
    add_export_options()

    # Add knowledge base
    add_knowledge_base()

if __name__ == "__main__":
    main()




    