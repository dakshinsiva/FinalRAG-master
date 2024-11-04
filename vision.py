import os
import logging
from dotenv import load_dotenv
from typing import List, Dict, Optional, Union
import datetime
from langchain_community.document_loaders import DirectoryLoader, PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain.chains import RetrievalQA
from langchain_openai import ChatOpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from functools import lru_cache

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Batch Processing Configuration
BATCH_SIZE = 10
MAX_BATCH_MB = 50
CONCURRENT_LIMIT = 5

# Cache Configuration
CACHE_SIZE_LIMIT = 1024 * 1024 * 1024  # 1GB
CACHE_TTL = 7 * 24 * 60 * 60  # 7 days

# Resource Configuration
# NUM_THREADS = min(16, os.cpu_count() * 2)
# CHUNK_SIZE = 800 if total_doc_size > 1_000_000 else 500
# CHUNK_OVERLAP = 100 if total_doc_size > 1_000_000 else 50

# Cost Optimization
MAX_TOKENS_PER_REQUEST = 4096
RESPONSE_TOKEN_LIMIT = 1000
BATCH_EMBEDDING_SIZE = 100  # documents per batch

class SecurityQuestionnaire:
    def __init__(self):
        self.sections = {
            "vulnerability_management": "Vulnerability Assessment and Management",
            "configuration_security": "Configuration and Security Controls",
            "database_security": "Database Security",
            "system_security": "System Security",
            "iam": "Identity and Access Management",
            "backup": "Backup and Recovery",
            "monitoring": "Monitoring and Logging",
            "infrastructure": "Infrastructure Security",
            "1B_questionnaire": "Security Assessment Questionnaire",
            "2a_vendor_selection": "Vendor Selection and Assessment",
            "3b_access_form": "Third Party Access Form"
        }
        self.questions = {
            "vulnerability_management": {
                "VA_PT_reports": "What is your vulnerability assessment and penetration testing process?",
                "remediation_status": "How do you track and manage remediation of security findings?",
                "security_reporting_sources": "What are your security reporting and monitoring sources?"
            },
            "configuration_security": {
                "ssl_certificate": "Describe your SSL/TLS certificate management process",
                "object_versioning": "How is object versioning implemented?",
                "encryption": "What encryption methods are used for data at rest and in transit?",
                "public_access": "How is public access to resources controlled?",
                "masking_mechanism": "What data masking mechanisms are in place?",
                "data_storage_location": "Where is sensitive data stored?",
                "key_storage": "How are encryption keys managed and stored?",
                "deletion_protection": "What deletion protection measures are in place?",
                "auto_healing": "Describe your auto-healing capabilities",
                "iam_configuration": "Detail your IAM configuration and policies"
            },
            "database_security": {
                "db_encryption": "What database encryption methods are used?",
                "db_access_control": "How is database access controlled?",
                "db_monitoring": "What database monitoring is in place?",
                "db_backup": "How are database backups managed?"
            },
            "system_security": {
                "os_hardening": "What OS hardening measures are implemented?",
                "patch_management": "Describe your patch management process",
                "endpoint_protection": "What endpoint protection is in place?",
                "system_monitoring": "How are systems monitored?"
            },
            "iam": {
                "access_control": "How is access control managed?",
                "authentication": "What authentication methods are used?",
                "privileged_access": "How is privileged access managed?",
                "user_lifecycle": "Describe the user lifecycle management process"
            },
            "backup": {
                "backup_frequency": "What is your backup frequency?",
                "backup_storage": "Where are backups stored?",
                "backup_testing": "How often are backups tested?",
                "recovery_process": "Describe your recovery process"
            },
            "monitoring": {
                "log_management": "How are logs managed and stored?",
                "alert_system": "What alerting systems are in place?",
                "incident_response": "Describe your incident response process",
                "monitoring_tools": "What monitoring tools are used?"
            },
            "infrastructure": {
                "network_security": "What network security measures are in place?",
                "cloud_security": "How is cloud security managed?",
                "physical_security": "What physical security controls exist?",
                "disaster_recovery": "Describe your disaster recovery plan"
            },
            "1B_questionnaire": {
                "outsourcing_nature": "What is the nature of process/service to be outsourced?",
                "spii_access": "Does the process/service/solution involve accessing or processing of SPII data of ABC's or Intellectual property (e.g. trade secrets, trademark, patents etc. ) details?",
                "spii_sharing": "Does the process/service/solution involve Sharing of SPII data of ABC's or Intellectual property (e.g. trade secrets, trademark, patents etc. ) details?",
                "spii_storage": "Does the process/service/solution involve storing of SPII data of ABC's or Intellectual property (e.g. trade secrets, trademark, patents etc. ) details?",
                "pii_access": "Does the process/service/solution involve accessing or processing of PII data of ABC customers or employees?",
                "pii_sharing": "Does the process/service/solution involve sharing of ABC customers' or employees' PII data?",
                "pii_storage": "Does the process/service/solution involve storing of ABC customers' or employees' PII data on the third party systems?",
                "business_critical_access": "Does the process/service/solution involve accessing or processing of ABC's business critical / confidential data (e.g. financial / Market research data)?",
                "business_critical_sharing": "Does the process/service/solution involve sharing of ABC's business critical / confidential data (e.g. Financial / Market research data)?",
                "business_critical_storage": "Does the process/service/solution involve storing of ABC's critical / confidential data (e.g. financial / Market research data) on third party systems?",
                "pii_records": "How many customers' or employees' PII/ SPI records will be accessed, processed, or shared for this process?",
                "data_origin": "At which location(s) will the data be originated?",
                "cross_border_transfer": "Does this process/service/solution involve cross-border transfer(sending data outside the country) of data?",
                "network_access": "Does the process/service/solution require access to ABC's internal network?",
                "network_access_level": "What level of network access will be required for this process/service/solution?",
                "remote_transfer": "Does the process/service/solution require transfer of ABC's confidential information or PII/SPI data via remote mechanism?",
                "business_disruption": "Would there be significant disruption to ABC's business operation if this process/service/solution stops abruptly?",
                "customer_impact": "Would the sudden failure of this process/service/solution impact ABC's customers?",
                "restoration_impact": "If the time required to restore the service/process is more than 24 hours, would there be a negative impact to ABC?",
                "brand_harm": "Would the process/service/solution failure cause significant harm to ABC's brand or reputation?",
                "revenue_impact": "Would the sudden failure of this process/service/solution impact ABC's revenue?",
                "regulatory_risk": "Would ABC be subjected to regulatory review, enforcement actions, or fines if there is a failure or interruption of the process/service/solution?",
                "local_regulations": "Is the process/service/solution subjected to local regulations / jurisdictions as a result of cross-border operations in multiple regions?"
            },
            "2a_vendor_selection": {
                "services_provided": "What services are provided by the vendor to ABC?",
                "data_access": "Which ABC data would you need access to (Personally identifiable information (PII) or SPII data or Business critical / confidential data (e.g. financial / business research data) or Intellectual property (e.g. trade secrets, trademark, patents etc)). Please select all the data classification you have access to",
                "data_storage": "Would you store any ABC data?",
                "payment_card_info": "Do you process/store payment card information data?",
                "data_storage_location_geo": "Where will you store the ABC data - Geographically?",
                "data_storage_location": "Where will you store the ABC data?",
                "iso_27001": "Is the vendor processing centre ISO 27001 (Information Security Management System- ISMS)?",
                "iso_22301": "Is the vendor processing centre ISO 22301 (Business continuity)?",
                "pci_dss": "Is the vendor processing centre PCI DSS compliant?",
                "soc_compliance": "Is the vendor processing centre SOC 1 or SOC 2 compliant?",
                "info_security_policy": "Does your organisation have Information Security Policy in line with industry best practices such as ISO27001 standard and has this been reviewed in the last 12 months?",
                "data_privacy_policy": "Does your organisation have data privacy policy and has this been reviewed in the last 12 months?",
                "security_team": "Do you have a Information security team?",
                "vulnerability_management": "Does your organisation have a Vulnerability Management and Penetration testing program that ensures that application vulnerabilities are addressed in a timely manner?",
                "logging_monitoring": "Do you have logging and monitoring enabled on all your information processing systems and maintain the logs securely for a rolling period of 180 days?",
                "security_operations": "Do you have a security operations centre and SIEM technology?",
                "endpoint_security": "Do you have end point security solutions implemented (Anti-virus and DLP)?",
                "threat_intelligence": "Do you have threat intelligence tools or controls implemented?",
                "application_security": "Do you perform application security assessments as per OWASP standards?",
                "network_monitoring": "Do you have network perimeter monitoring devices such as IDS/ IPS/ Next generation firewall/ application firewall/ proxy?",
                "usage_policy": "Do you have an acceptable usage policy which restricts printing of confidential information?",
                "data_retention_policy": "Do you have a data retention and disposal policy?",
                "idam_solution": "Do you have an Identity access management solution (IDAM) for normal user IDs?",
                "pim_pam_solution": "Do you have a privileged identity access management solution (PIM/ PAM) for privileged user IDs?",
                "user_access_review": "Do you perform a user access review for normal and privileged user IDs?",
                "change_management": "Do you have a change management policy in place?",
                "backup_restoration": "Do you have a data backup and restoration process?",
                "data_encryption": "Do you use data encryption/ cryptography?",
                "cyber_crisis_management": "Do you have a cyber crisis notification and management process?",
                "patch_management": "Do you have patch management process?",
                "breach_history": "Please mark no if you have been breached in last 12 months?",
                "physical_security": "Do you have physical security and environmental controls such as CCTV monitoring equipment security, physical ID and access cards, etc.?",
                "cloud_security_policy": "Do you have a cloud security policy?",
                "data_controls": "Do you have controls implemented for data at rest and data in transit?",
                "security_awareness": "Do you have a security awareness program for your employees, part-timers?",
                "cmdb_asset_discovery": "Do you use CMDB and asset discovery tool?",
                "business_continuity_policy": "Is there a Business Continuity policy, IT Disaster Recovery policy (and supporting documentation) which maintained, reviewed and signed off in the last 12 months?"
            },
            "3b_access_form": {
                "product_service_description": "Describe the product or service to be provided",
                "third_party_access_type": "Classify the type of third party access:\n(Service Provider\nManaged Services\nCustomer Access\nOutsourcing\nAuditor\nConsultant\nDeveloper\nCleaning/Catering/Support services\nTemp/Work Experience/Student Placement)",
                "systems_applications_access": "What systems or applications are to be accessed logically – if any?",
                "site_access": "Will the third party require access to the ABC site?",
                "computer_room_access": "Will the third party require access to the ABC computer room?",
                "remote_access": "Will the third party require Remote Access – if so describe how this will be achieved?",
                "sensitive_info_access": "Will the third party require access to any sensitive or confidential information – if so describe.",
                "personnel_count": "How many personnel from the third party will require access?",
                "third_party_contacts": "Who is the/are the key contacts at the third party (provide email addresses/telephone numbers)?",
                "sponsor_contact": "Which Aggregate employee or contact will sponsor the access?",
                "access_duration": "How long is the access required for?",
                "asset_requirements": "Does the third party require any Aggregate owned assets?",
                "service_continuity_impact": "Does this third party access have any impact on IT service continuity?",
                "it_stakeholders": "Identify any key contacts or stakeholders in IT related to this third party access.",
                "business_stakeholders": "Identify any key contacts or stakeholders in the business related to this third party access.",
                "legal_regulatory_impact": "Is there any identified legal or regulatory impact?",
                "contract_status": "Is there a valid and signed contract or MSA signed with the vendor organization?",
                "nda_status": "Are the NDA and Confidentiality Agreements signed by the vendor?"
            }
        }

    def get_section_name(self, section_key):
        return self.sections.get(section_key, section_key)

    def get_questions_by_section(self, section_key):
        return self.questions.get(section_key, {})

@lru_cache(maxsize=1000)
def get_embedding(text: str) -> List[float]:
    return embeddings.embed_query(text)

def write_formatted_results(results, questionnaire):
    """Write formatted results to both text and Word files"""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H-%M-%S")
    txt_output = f"security_questionnaire_responses_{timestamp}.txt"
    docx_output = f"security_questionnaire_responses_{timestamp}.docx"
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    title = doc.add_heading('Security Questionnaire Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    timestamp_para = doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    with open(txt_output, 'w', encoding='utf-8') as f:
        f.write("╔═══════════════════════════════════════════════════════════════════════════╗\n")
        f.write("║                      Security Questionnaire Analysis                         ║\n")
        f.write("║                                                                             ║\n")
        f.write(f"║  Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}                                     ║\n")
        f.write("╚═══════════════════════════════════════════════════════════════════════════════╝\n\n")

        for section, questions in results['answers'].items():
            section_name = questionnaire.sections.get(section, section)
            
            f.write("\n" + "=" * 80 + "\n")
            f.write(f"{section_name.upper()}\n")
            f.write("=" * 80 + "\n\n")
            
            doc.add_heading(section_name.upper(), level=1)
            
            for q_key, answer in questions.items():
                full_question = questionnaire.questions[section][q_key]
                
                f.write("📝 Question:\n")
                f.write(f"{full_question}\n\n")
                f.write("🔍 Answer:\n")
                f.write(f"{answer['result'] if isinstance(answer, dict) else answer}\n\n")
                
                q_para = doc.add_paragraph()
                q_para.add_run("Question:\n").bold = True
                q_para.add_run(full_question)
                
                a_para = doc.add_paragraph()
                a_para.add_run("Answer:\n").bold = True
                a_para.add_run(answer['result'] if isinstance(answer, dict) else answer)
                
                source_docs = []
                if isinstance(answer, dict) and 'source_documents' in answer:
                    source_docs = answer['source_documents']
                
                f.write("📚 Source Documents:\n")
                if source_docs:
                    for doc_ref in source_docs:
                        source = os.path.basename(doc_ref.metadata.get('source', '[Document reference]'))
                        page_num = doc_ref.metadata.get('page', 1)
                        f.write(f"• {source} (Page {page_num})\n")
                else:
                    f.write("• [Document reference]\n")
                
                s_para = doc.add_paragraph()
                s_para.add_run("Source Documents:\n").bold = True
                if source_docs:
                    for doc_ref in source_docs:
                        source = os.path.basename(doc_ref.metadata.get('source', '[Document reference]'))
                        page_num = doc_ref.metadata.get('page', 1)
                        s_para.add_run(f"• {source} (Page {page_num})\n")
                else:
                    s_para.add_run("• [Document reference]")
                
                f.write("\n" + "_" * 80 + "\n\n")
                doc.add_paragraph("_" * 80)
                doc.add_paragraph()

    doc.save(docx_output)
    return txt_output, docx_output

def main():
    """Main function for RAG system"""
    try:
        docs_path = "/Users/dakshinsiva/final_RAG/docs"
        documents = []
        
        loaders = {
            '.pdf': PyPDFLoader,
        }
        
        for filename in os.listdir(docs_path):
            file_path = os.path.join(docs_path, filename)
            file_extension = os.path.splitext(filename)[1].lower()
            
            if file_extension in loaders:
                try:
                    loader = loaders[file_extension](file_path)
                    docs = loader.load()
                    documents.extend(docs)
                    logger.info(f"Successfully loaded {filename}")
                except Exception as e:
                    logger.error(f"Error loading file {filename}: {str(e)}")
                    continue
        
        if not documents:
            raise ValueError("No documents were successfully loaded")
            
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=100,
            add_start_index=True,
            separators=["\n\n", "\n", " ", ""]
        )
        texts = text_splitter.split_documents(documents)
        
        vector_store = FAISS.from_documents(
            texts, 
            OpenAIEmbeddings()
        )
        
        retriever = vector_store.as_retriever(
            search_type="similarity", 
            search_kwargs={
                "k": 3,
                "include_metadata": True,
                "score_threshold": 0.7
            }
        )
        
        llm = ChatOpenAI(temperature=0)
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="refine",
            retriever=retriever, # or "map_reduce", "refine", "map_rerank"
            return_source_documents=True,
            verbose=True
        )
        
        questionnaire = SecurityQuestionnaire()
        
        results = {
            'questions': questionnaire.questions,
            'sections': questionnaire.sections,
            'answers': {}
        }
        
        for section, questions in questionnaire.questions.items():
            results['answers'][section] = {}
            for key, question in questions.items():
                try:
                    answer = qa_chain({"query": question})
                    results['answers'][section][key] = answer
                except Exception as e:
                    logger.error(f"Error processing question {key} in section {section}: {str(e)}")
                    results['answers'][section][key] = {"error": str(e)}
        
        txt_file, docx_file = write_formatted_results(results, questionnaire)
        
        print(f"\nAnalysis complete!")
        print(f"Results saved to:")
        print(f"- Text file: {txt_file}")
        print(f"- Word document: {docx_file}")
        
    except Exception as e:
        logger.error(f"Main execution failed: {str(e)}")
        print(f"Error: {str(e)}")

if __name__ == "__main__":
    main()
